from docx import Document
import xml.etree.ElementTree as ET

def extract_formulas_from_docx(doc_path):
    doc = Document(doc_path)
    
    print("=" * 70)
    print(f"分析文档: {doc_path}")
    print("=" * 70)
    print()
    
    print(f"文档共有 {len(doc.paragraphs)} 个段落")
    print()
    
    formulas_found = []
    paragraph_count = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        paragraph_count += 1
        
        if text:
            print(f"段落 {i} [{para.style.name}]: {text[:200]}")
            print()
        
        math_elements = []
        
        for run in para.runs:
            if hasattr(run._element, 'getiterator'):
                for element in run._element.getiterator():
                    tag = element.tag
                    if 'm:oMath' in tag or 'math' in tag.lower() or 'Math' in tag:
                        math_elements.append(element)
            elif hasattr(run._element, 'iter'):
                for element in run._element.iter():
                    tag = element.tag
                    if 'm:oMath' in tag or 'math' in tag.lower() or 'Math' in tag:
                        math_elements.append(element)
        
        if math_elements:
            print(f"  ↳ 在段落 {i} 中找到 {len(math_elements)} 个公式元素")
            for idx, elem in enumerate(math_elements):
                print(f"    - 公式 {idx + 1}, 标签: {elem.tag}")
                formulas_found.append((i, elem))
    
    print()
    print("=" * 70)
    print(f"统计: 共找到 {len(formulas_found)} 个公式")
    print("=" * 70)
    
    return doc, formulas_found

def inspect_document_structure(doc):
    print()
    print("文档结构检查:")
    print(f"  段落数: {len(doc.paragraphs)}")
    print(f"  文档有 {len(doc.sections)} 个章节")
    
    for idx, table in enumerate(doc.tables):
        print(f"  表格 {idx+1}: {len(table.rows)} 行, {len(table.columns)} 列")
    
    return doc

if __name__ == "__main__":
    doc_path = r'c:\Users\licl45\Desktop\testTRAE\testword.docx'
    doc, formulas = extract_formulas_from_docx(doc_path)
    inspect_document_structure(doc)