"""
编辑 Word 文档：插入丰富的公式和内容
"""
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def insert_formula_simple(doc, formula_xml_str, after_paragraph_idx=None):
    """使用 lxml 直接操作 XML 插入 OMML 公式"""
    if after_paragraph_idx is not None and after_paragraph_idx < len(doc.paragraphs):
        para_element = doc.paragraphs[after_paragraph_idx]._element
        new_p = etree.Element(qn('w:p'))
        para_element.addnext(new_p)
    else:
        new_p = etree.SubElement(doc.element.body, qn('w:p'))
    
    new_r = etree.SubElement(new_p, qn('w:r'))
    rPr = etree.SubElement(new_r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Cambria Math')
    rFonts.set(qn('w:hAnsi'), 'Cambria Math')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '24')
    
    omath_para = etree.SubElement(new_r, qn('m:oMathPara'))
    formula_element = etree.fromstring(formula_xml_str)
    omath_para.append(formula_element)
    
    return new_p

def build_rich_word_document(input_path, output_path):
    doc = Document(input_path)
    
    # 添加分隔标题
    doc.add_heading('丰富的数学公式集合', level=1)
    
    # ======== 公式 3: 欧拉公式 ========
    doc.add_heading('欧拉公式 (Euler\'s Formula)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('欧拉公式被誉为"最美的数学公式"，它将自然常数 e、虚数单位 i、圆周率 π 联系在一起。').font.size = Pt(11)
    
    formula_euler = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>e</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e><m:r><m:t>i</m:t></m:r><m:r><m:t>θ</m:t></m:r></m:e>
            <m:sup><m:r><m:t>i</m:t></m:r><m:r><m:t>θ</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>cos</m:t></m:r>
        <m:r><m:t>θ</m:t></m:r>
        <m:r><m:t>+</m:t></m:r>
        <m:r><m:t>i</m:t></m:r>
        <m:r><m:t>sin</m:t></m:r>
        <m:r><m:t>θ</m:t></m:r>
    </m:oMath>'''
    insert_formula_simple(doc, formula_euler, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    p = doc.add_paragraph()
    p.add_run('当 θ = π 时，得到最著名的形式：').font.size = Pt(11)
    
    formula_euler2 = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>e</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e><m:r><m:t>i</m:t></m:r><m:r><m:t>π</m:t></m:r></m:e>
            <m:sup><m:r><m:t>i</m:t></m:r><m:r><m:t>π</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>+</m:t></m:r>
        <m:r><m:t>1</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>0</m:t></m:r>
    </m:oMath>'''
    insert_formula_simple(doc, formula_euler2, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ======== 公式 4: 正态分布 ========
    doc.add_heading('正态分布 (Normal Distribution)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('正态分布是统计学中最重要的概率分布，也称为高斯分布。').font.size = Pt(11)
    
    formula_normal = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>f</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>x</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:f>
            <m:fPr/>
            <m:num><m:r><m:t>1</m:t></m:r></m:num>
            <m:den>
                <m:r><m:t>σ</m:t></m:r>
                <m:rad>
                    <m:radPr/>
                    <m:e><m:r><m:t>2</m:t></m:r><m:r><m:t>π</m:t></m:r></m:e>
                </m:rad>
            </m:den>
        </m:f>
        <m:r><m:t>e</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup>
                <m:r><m:t>-</m:t></m:r>
                <m:f>
                    <m:fPr/>
                    <m:num>
                        <m:r><m:t>(</m:t></m:r>
                        <m:r><m:t>x</m:t></m:r>
                        <m:r><m:t>-</m:t></m:r>
                        <m:r><m:t>μ</m:t></m:r>
                        <m:r><m:t>)</m:t></m:r>
                        <m:sSup>
                            <m:sSupPr/>
                            <m:e/>
                            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
                        </m:sSup>
                    </m:num>
                    <m:den>
                        <m:r><m:t>2</m:t></m:r>
                        <m:r><m:t>σ</m:t></m:r>
                        <m:sSup>
                            <m:sSupPr/>
                            <m:e/>
                            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
                        </m:sSup>
                    </m:den>
                </m:f>
            </m:sup>
        </m:sSup>
    </m:oMath>'''
    insert_formula_simple(doc, formula_normal, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    p = doc.add_paragraph()
    p.add_run('其中 μ 是均值，σ 是标准差。').font.size = Pt(11)
    
    # ======== 公式 5: 勾股定理 ========
    doc.add_heading('勾股定理 (Pythagorean Theorem)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('勾股定理是几何学中最基本的定理之一，描述了直角三角形三边之间的关系。').font.size = Pt(11)
    
    formula_pythag = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>a</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>+</m:t></m:r>
        <m:r><m:t>b</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>c</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
    </m:oMath>'''
    insert_formula_simple(doc, formula_pythag, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ======== 公式 6: 贝叶斯定理 ========
    doc.add_heading('贝叶斯定理 (Bayes\' Theorem)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('贝叶斯定理是概率论中的核心定理，描述了条件概率之间的关系。').font.size = Pt(11)
    
    formula_bayes = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>P</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>A</m:t></m:r>
        <m:r><m:t>|</m:t></m:r>
        <m:r><m:t>B</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:f>
            <m:fPr/>
            <m:num>
                <m:r><m:t>P</m:t></m:r>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>B</m:t></m:r>
                <m:r><m:t>|</m:t></m:r>
                <m:r><m:t>A</m:t></m:r>
                <m:r><m:t>)</m:t></m:r>
                <m:r><m:t>P</m:t></m:r>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>A</m:t></m:r>
                <m:r><m:t>)</m:t></m:r>
            </m:num>
            <m:den>
                <m:r><m:t>P</m:t></m:r>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>B</m:t></m:r>
                <m:r><m:t>)</m:t></m:r>
            </m:den>
        </m:f>
    </m:oMath>'''
    insert_formula_simple(doc, formula_bayes, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ======== 公式 7: 定积分 ========
    doc.add_heading('积分公式 (Integral Formula)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('定积分是微积分中的基本概念，表示函数曲线下的面积。').font.size = Pt(11)
    
    formula_integral = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:nary>
            <m:naryPr>
                <m:chr m:val="∫"/>
                <m:limLoc m:val="subSup"/>
            </m:naryPr>
            <m:sub><m:r><m:t>a</m:t></m:r></m:sub>
            <m:sup><m:r><m:t>b</m:t></m:r></m:sup>
            <m:e>
                <m:r><m:t>f</m:t></m:r>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>x</m:t></m:r>
                <m:r><m:t>)</m:t></m:r>
                <m:r><m:t>d</m:t></m:r>
                <m:r><m:t>x</m:t></m:r>
            </m:e>
        </m:nary>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>F</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>b</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
        <m:r><m:t>-</m:t></m:r>
        <m:r><m:t>F</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>a</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
    </m:oMath>'''
    insert_formula_simple(doc, formula_integral, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ======== 公式 8: 泰勒级数 ========
    doc.add_heading('泰勒级数 (Taylor Series)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('泰勒级数将函数展开为无穷级数，是微积分中最重要的工具之一。').font.size = Pt(11)
    
    formula_taylor = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>f</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>x</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:nary>
            <m:naryPr>
                <m:chr m:val="∑"/>
                <m:limLoc m:val="subSup"/>
            </m:naryPr>
            <m:sub>
                <m:r><m:t>n</m:t></m:r>
                <m:r><m:t>=</m:t></m:r>
                <m:r><m:t>0</m:t></m:r>
            </m:sub>
            <m:sup><m:r><m:t>∞</m:t></m:r></m:sup>
            <m:e>
                <m:f>
                    <m:fPr/>
                    <m:num>
                        <m:r><m:t>f</m:t></m:r>
                        <m:sSup>
                            <m:sSupPr/>
                            <m:e/>
                            <m:sup>
                                <m:r><m:t>(</m:t></m:r>
                                <m:r><m:t>n</m:t></m:r>
                                <m:r><m:t>)</m:t></m:r>
                            </m:sup>
                        </m:sSup>
                        <m:r><m:t>(</m:t></m:r>
                        <m:r><m:t>a</m:t></m:r>
                        <m:r><m:t>)</m:t></m:r>
                    </m:num>
                    <m:den>
                        <m:r><m:t>n</m:t></m:r>
                        <m:r><m:t>!</m:t></m:r>
                    </m:den>
                </m:f>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>x</m:t></m:r>
                <m:r><m:t>-</m:t></m:r>
                <m:r><m:t>a</m:t></m:r>
                <m:sSup>
                    <m:sSupPr/>
                    <m:e><m:r><m:t>)</m:t></m:r></m:e>
                    <m:sup><m:r><m:t>n</m:t></m:r></m:sup>
                </m:sSup>
            </m:e>
        </m:nary>
    </m:oMath>'''
    insert_formula_simple(doc, formula_taylor, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ======== 公式 9: 麦克斯韦方程组 ========
    doc.add_heading('麦克斯韦方程组 (Maxwell\'s Equations)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('麦克斯韦方程组是电磁学的基本方程，描述了电场和磁场的关系。').font.size = Pt(11)
    
    formula_maxwell1 = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:nary>
            <m:naryPr>
                <m:chr m:val="∮"/>
                <m:limLoc m:val="subSup"/>
            </m:naryPr>
            <m:sub><m:r><m:t>S</m:t></m:r></m:sub>
            <m:sup/>
            <m:e>
                <m:r><m:t>E</m:t></m:r>
                <m:r><m:t>·</m:t></m:r>
                <m:r><m:t>d</m:t></m:r>
                <m:r><m:t>A</m:t></m:r>
            </m:e>
        </m:nary>
        <m:r><m:t>=</m:t></m:r>
        <m:f>
            <m:fPr/>
            <m:num><m:r><m:t>q</m:t></m:r></m:num>
            <m:den><m:r><m:t>ε</m:t></m:r><m:sSub><m:sSubPr/><m:e/><m:sub><m:r><m:t>0</m:t></m:r></m:sub></m:sSub></m:den>
        </m:f>
    </m:oMath>'''
    insert_formula_simple(doc, formula_maxwell1, after_paragraph_idx=len(doc.paragraphs) - 1)
    p = doc.add_paragraph()
    p.add_run('高斯电场定律 (Gauss\'s Law for Electricity)').font.size = Pt(10)
    
    formula_maxwell2 = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:nary>
            <m:naryPr>
                <m:chr m:val="∮"/>
                <m:limLoc m:val="subSup"/>
            </m:naryPr>
            <m:sub><m:r><m:t>S</m:t></m:r></m:sub>
            <m:sup/>
            <m:e>
                <m:r><m:t>B</m:t></m:r>
                <m:r><m:t>·</m:t></m:r>
                <m:r><m:t>d</m:t></m:r>
                <m:r><m:t>A</m:t></m:r>
            </m:e>
        </m:nary>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>0</m:t></m:r>
    </m:oMath>'''
    insert_formula_simple(doc, formula_maxwell2, after_paragraph_idx=len(doc.paragraphs) - 1)
    p = doc.add_paragraph()
    p.add_run('高斯磁场定律 (Gauss\'s Law for Magnetism)').font.size = Pt(10)
    
    # ======== 公式 10: 质能方程 ========
    doc.add_heading('质能方程 (Mass-Energy Equivalence)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('爱因斯坦的质能方程是物理学中最著名的公式之一，揭示了质量和能量的等价关系。').font.size = Pt(11)
    
    formula_einstein = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>E</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>m</m:t></m:r>
        <m:r><m:t>c</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
    </m:oMath>'''
    insert_formula_simple(doc, formula_einstein, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    p = doc.add_paragraph()
    p.add_run('其中 E 是能量，m 是质量，c 是光速 (≈ 3×10⁸ m/s)。').font.size = Pt(11)
    
    # ======== 公式总结 ========
    doc.add_heading('公式总结', level=1)
    
    p = doc.add_paragraph()
    p.add_run('本文档包含以下公式：').font.size = Pt(11)
    
    formulas_list = [
        '1. 一元二次方程求根公式: x = (-b ± √(b² - 4ac)) / 2a',
        '2. 傅里叶级数: f(x) = a₀ + Σ(aₙcos(nπx/L) + bₙsin(nπx/L))',
        '3. 欧拉公式: e^(iθ) = cosθ + i·sinθ',
        '4. 正态分布: f(x) = 1/(σ√(2π)) · e^(-(x-μ)²/(2σ²))',
        '5. 勾股定理: a² + b² = c²',
        '6. 贝叶斯定理: P(A|B) = P(B|A)·P(A) / P(B)',
        '7. 定积分: ∫[a,b] f(x)dx = F(b) - F(a)',
        '8. 泰勒级数: f(x) = Σ f^(n)(a)/n! · (x-a)^n',
        '9. 麦克斯韦方程组',
        '10. 质能方程: E = mc²',
    ]
    
    for formula in formulas_list:
        p = doc.add_paragraph()
        p.add_run('• ' + formula).font.size = Pt(11)
    
    doc.save(output_path)
    print(f"丰富的 Word 文档已保存到: {output_path}")
    print(f"共添加了 8 个新公式，总计 10 个公式")

if __name__ == "__main__":
    input_path = r'c:\Users\licl45\Desktop\testTRAE\testword.docx'
    output_path = r'c:\Users\licl45\Desktop\testTRAE\testword_enriched.docx'
    build_rich_word_document(input_path, output_path)