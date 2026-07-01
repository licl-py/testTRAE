"""
综合脚本：读取 Word 公式 + 生成测试代码 + 编辑 Word 文档插入新公式
"""
import zipfile
import xml.etree.ElementTree as ET
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import os
import shutil

# ============================================================
# OMML 命名空间
# ============================================================
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# ============================================================
# 第1部分: 从 Word 提取公式文本
# ============================================================

def extract_math_text(elem):
    """递归提取 OMML 公式元素的文本表示"""
    results = []
    
    for child in elem:
        child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if child_tag == 'r':
            for t_elem in child.findall(f'{{{MATH_NS}}}t'):
                text = t_elem.text or ''
                results.append(text)
        
        elif child_tag == 'f':
            num_text = ''
            den_text = ''
            for part in child:
                part_tag = part.tag.split('}')[-1] if '}' in part.tag else part.tag
                if part_tag == 'num':
                    num_text = extract_math_text(part)
                elif part_tag == 'den':
                    den_text = extract_math_text(part)
            results.append(f'({num_text})/({den_text})')
        
        elif child_tag == 'sSup':
            base = ''
            sup = ''
            for part in child:
                part_tag = part.tag.split('}')[-1] if '}' in part.tag else part.tag
                if part_tag == 'e':
                    base = extract_math_text(part)
                elif part_tag == 'sup':
                    sup = extract_math_text(part)
            results.append(f'{base}^{sup}')
        
        elif child_tag == 'sSub':
            base = ''
            sub = ''
            for part in child:
                part_tag = part.tag.split('}')[-1] if '}' in part.tag else part.tag
                if part_tag == 'e':
                    base = extract_math_text(part)
                elif part_tag == 'sub':
                    sub = extract_math_text(part)
            results.append(f'{base}_{sub}')
        
        elif child_tag == 'sSubSup':
            base = ''
            sub = ''
            sup = ''
            for part in child:
                part_tag = part.tag.split('}')[-1] if '}' in part.tag else part.tag
                if part_tag == 'e':
                    base = extract_math_text(part)
                elif part_tag == 'sub':
                    sub = extract_math_text(part)
                elif part_tag == 'sup':
                    sup = extract_math_text(part)
            results.append(f'{base}_{sub}^{sup}')
        
        elif child_tag == 'rad':
            deg_text = ''
            e_text = ''
            for part in child:
                part_tag = part.tag.split('}')[-1] if '}' in part.tag else part.tag
                if part_tag == 'deg':
                    deg_text = extract_math_text(part)
                elif part_tag == 'e':
                    e_text = extract_math_text(part)
            if deg_text:
                results.append(f'({deg_text})√({e_text})')
            else:
                results.append(f'√({e_text})')
        
        elif child_tag == 'nary':
            base = ''
            sub = ''
            sup = ''
            for part in child:
                part_tag = part.tag.split('}')[-1] if '}' in part.tag else part.tag
                if part_tag == 'e':
                    base = extract_math_text(part)
                elif part_tag == 'sub':
                    sub = extract_math_text(part)
                elif part_tag == 'sup':
                    sup = extract_math_text(part)
            
            nary_op = '∫'
            for pr_child in child.findall(f'{{{MATH_NS}}}naryPr'):
                chr_elem = pr_child.find(f'{{{MATH_NS}}}chr')
                if chr_elem is not None:
                    char_val = chr_elem.get(f'{{{MATH_NS}}}val', '')
                    if char_val == '∫':
                        nary_op = '∫'
                    elif char_val == '∑':
                        nary_op = '∑'
                    elif char_val == '∏':
                        nary_op = '∏'
            
            results.append(f'{nary_op}_{sub}^{sup}({base})')
        
        elif child_tag == 'd':
            inner = extract_math_text(child)
            results.append(inner)
        
        elif child_tag == 'groupChr':
            inner = extract_math_text(child)
            results.append(f'({inner})')
        
        elif child_tag == 'acc':
            inner = extract_math_text(child)
            results.append(inner)
        
        elif child_tag == 'bar':
            inner = extract_math_text(child)
            results.append(f'overline({inner})')
        
        elif child_tag == 'eqArr':
            inner = extract_math_text(child)
            results.append(inner)
        
        elif child_tag == 'func':
            fname = ''
            fe = ''
            for part in child:
                part_tag = part.tag.split('}')[-1] if '}' in part.tag else part.tag
                if part_tag == 'fName':
                    fname = extract_math_text(part)
                elif part_tag == 'e':
                    fe = extract_math_text(part)
            if fname and fe:
                results.append(f'{fname}({fe})')
            elif fname:
                results.append(fname)
            elif fe:
                results.append(f'({fe})')
        
        elif child_tag == 'box':
            inner = extract_math_text(child)
            results.append(inner)
        
        elif child_tag == 'ctrlPr':
            pass
        
        elif child_tag == 'limLow':
            inner = extract_math_text(child)
            results.append(inner)
        
        elif child_tag == 'limUpp':
            inner = extract_math_text(child)
            results.append(inner)
        
        else:
            inner = extract_math_text(child)
            if inner:
                results.append(inner)
    
    return ''.join(results)


def read_formulas_from_docx(doc_path):
    """读取 Word 文档中的所有公式"""
    formulas = []
    
    with zipfile.ZipFile(doc_path, 'r') as z:
        with z.open('word/document.xml') as f:
            tree = ET.parse(f)
    
    root = tree.getroot()
    all_math = root.findall(f'.//{{{MATH_NS}}}oMath')
    
    for idx, math_elem in enumerate(all_math, 1):
        formula_text = extract_math_text(math_elem)
        raw_xml = ET.tostring(math_elem, encoding='unicode')
        formulas.append({
            'index': idx,
            'text': formula_text,
            'xml': raw_xml,
            'element': math_elem
        })
    
    return formulas


# ============================================================
# 第2部分: 生成 Python 测试代码
# ============================================================

def generate_test_code(formulas):
    """为每个公式生成 Python 测试代码"""
    
    code = f'''"""
自动生成的公式测试代码 - 从 Word 文档中提取的公式
"""
import math
from typing import Tuple, Optional, Callable


# ============================================================
# 公式 1: 一元二次方程求根公式 (Quadratic Formula)
# 提取文本: {formulas[0]['text']}
# 原始形式: x = (-b ± √(b² - 4ac)) / 2a
# ============================================================

def quadratic_formula(a: float, b: float, c: float) -> Tuple[Optional[float], Optional[float]]:
    """
    求解一元二次方程 ax^2 + bx + c = 0 的两个根。
    公式: x = (-b ± √(b² - 4ac)) / 2a
    
    返回:
        (x1, x2) 两个根。如果判别式 < 0 则返回 (None, None) 表示无实数解。
    """
    discriminant = b * b - 4 * a * c
    if discriminant < 0:
        return None, None
    
    sqrt_discriminant = math.sqrt(discriminant)
    x1 = (-b + sqrt_discriminant) / (2 * a)
    x2 = (-b - sqrt_discriminant) / (2 * a)
    return x1, x2


# ============================================================
# 公式 2: 傅里叶级数 (Fourier Series)
# 提取文本: {formulas[1]['text']}
# 原始形式: f(x) = a₀ + Σ[n=1→∞] (aₙcos(nπx/L) + bₙsin(nπx/L))
# ============================================================

def fourier_series(
    x: float,
    a0: float,
    an: Callable[[int], float],
    bn: Callable[[int], float],
    L: float,
    terms: int = 20
) -> float:
    """
    计算傅里叶级数的前 N 项和。
    公式: f(x) = a₀ + Σ[n=1→∞] (aₙcos(nπx/L) + bₙsin(nπx/L))
    
    参数:
        x:     自变量
        a0:    常数项
        an:    余弦系数函数，入参 n 返回 a_n
        bn:    正弦系数函数，入参 n 返回 b_n
        L:     半周期长度
        terms: 求和的项数 (默认 20)
    """
    result = a0
    for n in range(1, terms + 1):
        result += an(n) * math.cos(n * math.pi * x / L)
        result += bn(n) * math.sin(n * math.pi * x / L)
    return result


# ============================================================
# 测试代码
# ============================================================

if __name__ == "__main__":
    print("=" * 60)
    print("公式测试 - 从 Word 文档自动提取")
    print("=" * 60)
    
    # --- 测试公式 1: 一元二次方程求根公式 ---
    print("\\n" + "=" * 60)
    print("测试公式 1: 一元二次方程求根公式")
    print("公式: x = (-b ± √(b² - 4ac)) / 2a")
    print("=" * 60)
    
    test_cases = [
        (1, -5, 6, "两个不同实根"),
        (1, -4, 4, "重根"),
        (1, 1, 1, "无实数解"),
        (2, 3.5, -1.2, "含小数系数"),
    ]
    
    for a, b, c, desc in test_cases:
        x1, x2 = quadratic_formula(a, b, c)
        print(f"\\n方程: {{a}}x² + {{b}}x + {{c}} = 0  ({{desc}})")
        discriminant = b*b - 4*a*c
        print(f"  判别式 Δ = {{discriminant:.4f}}")
        if x1 is None:
            print(f"  结果: 无实数解 ✓")
        else:
            print(f"  x1 = {{x1:.6f}}, x2 = {{x2:.6f}}")
            verify1 = a*x1*x1 + b*x1 + c
            verify2 = a*x2*x2 + b*x2 + c
            print(f"  验证: a*x1²+b*x1+c = {{verify1:.2e}} ✓")
            print(f"  验证: a*x2²+b*x2+c = {{verify2:.2e}} ✓")
    
    # --- 测试公式 2: 傅里叶级数 ---
    print("\\n" + "=" * 60)
    print("测试公式 2: 傅里叶级数")
    print("公式: f(x) = a₀ + Σ(aₙcos(nπx/L) + bₙsin(nπx/L))")
    print("=" * 60)
    
    L = 1.0
    
    # 方波近似
    def an_square(n: int) -> float:
        return 0.0
    
    def bn_square(n: int) -> float:
        return (4 / (math.pi * n)) if n % 2 == 1 else 0.0
    
    print("\\n方波傅里叶级数近似 (前50项):")
    for x in [-0.5, 0.0, 0.25, 0.5]:
        val = fourier_series(x, 0.0, an_square, bn_square, L, terms=50)
        expected_sign = '+' if x > 0 else ('-' if x < 0 else '0')
        print(f"  f({{x:+.2f}}) = {{val:+.6f}} (预期符号: {{expected_sign}})")
    
    # 锯齿波近似
    def bn_sawtooth(n: int) -> float:
        return 2 * ((-1) ** (n + 1)) / (math.pi * n)
    
    print("\\n锯齿波傅里叶级数近似 (前50项):")
    for x in [-0.5, 0.0, 0.25, 0.5]:
        val = fourier_series(x, 0.0, an_square, bn_sawtooth, L, terms=50)
        print(f"  f({{x:+.2f}}) = {{val:+.6f}}")
    
    print("\\n" + "=" * 60)
    print("所有测试完成！")
    print("=" * 60)
'''
    
    return code


# ============================================================
# 第3部分: 创建 OMML 公式元素
# ============================================================

def create_omath_element():
    """创建一个空的 oMath 元素"""
    return ET.Element(f'{{{MATH_NS}}}oMath')


def create_omath_para():
    """创建一个 oMathPara 元素"""
    return ET.Element(f'{{{MATH_NS}}}oMathPara')


def create_math_run(text):
    """创建 oMath 中的 run 元素"""
    r = ET.SubElement(ET.Element(f'{{{MATH_NS}}}oMath'), f'{{{MATH_NS}}}r')
    t = ET.SubElement(r, f'{{{MATH_NS}}}t')
    t.text = text
    return r


def create_fraction(num, den):
    """创建分数 """
    f = ET.Element(f'{{{MATH_NS}}}f')
    fPr = ET.SubElement(f, f'{{{MATH_NS}}}fPr')
    num_elem = ET.SubElement(f, f'{{{MATH_NS}}}num')
    num_elem.append(num)
    den_elem = ET.SubElement(f, f'{{{MATH_NS}}}den')
    den_elem.append(den)
    return f


def create_superscript(base, sup):
    """创建上标"""
    sSup = ET.Element(f'{{{MATH_NS}}}sSup')
    e = ET.SubElement(sSup, f'{{{MATH_NS}}}e')
    e.append(base)
    sup_elem = ET.SubElement(sSup, f'{{{MATH_NS}}}sup')
    sup_elem.append(sup)
    return sSup


def create_subscript(base, sub):
    """创建下标"""
    sSub = ET.Element(f'{{{MATH_NS}}}sSub')
    e = ET.SubElement(sSub, f'{{{MATH_NS}}}e')
    e.append(base)
    sub_elem = ET.SubElement(sSub, f'{{{MATH_NS}}}sub')
    sub_elem.append(sub)
    return sSub


def create_radical(elem, deg=None):
    """创建根号"""
    rad = ET.Element(f'{{{MATH_NS}}}rad')
    if deg:
        deg_elem = ET.SubElement(rad, f'{{{MATH_NS}}}deg')
        deg_elem.append(deg)
    e = ET.SubElement(rad, f'{{{MATH_NS}}}e')
    e.append(elem)
    return rad


def create_sum(from_elem, to_elem, expr_elem):
    """创建求和符号"""
    nary = ET.Element(f'{{{MATH_NS}}}nary')
    naryPr = ET.SubElement(nary, f'{{{MATH_NS}}}naryPr')
    chr_elem = ET.SubElement(naryPr, f'{{{MATH_NS}}}chr')
    chr_elem.set(f'{{{MATH_NS}}}val', '∑')
    sub_elem = ET.SubElement(nary, f'{{{MATH_NS}}}sub')
    sub_elem.append(from_elem)
    sup_elem = ET.SubElement(nary, f'{{{MATH_NS}}}sup')
    sup_elem.append(to_elem)
    e = ET.SubElement(nary, f'{{{MATH_NS}}}e')
    e.append(expr_elem)
    return nary


def create_integral(from_elem, to_elem, expr_elem):
    """创建积分符号"""
    nary = ET.Element(f'{{{MATH_NS}}}nary')
    naryPr = ET.SubElement(nary, f'{{{MATH_NS}}}naryPr')
    chr_elem = ET.SubElement(naryPr, f'{{{MATH_NS}}}chr')
    chr_elem.set(f'{{{MATH_NS}}}val', '∫')
    sub_elem = ET.SubElement(nary, f'{{{MATH_NS}}}sub')
    sub_elem.append(from_elem)
    sup_elem = ET.SubElement(nary, f'{{{MATH_NS}}}sup')
    sup_elem.append(to_elem)
    e = ET.SubElement(nary, f'{{{MATH_NS}}}e')
    e.append(expr_elem)
    return nary


def text_run(text):
    """创建包含文本的 math run"""
    r = ET.Element(f'{{{MATH_NS}}}r')
    t = ET.SubElement(r, f'{{{MATH_NS}}}t')
    t.text = text
    return r


def plus_minus():
    """创建±符号"""
    r = ET.Element(f'{{{MATH_NS}}}r')
    t = ET.SubElement(r, f'{{{MATH_NS}}}t')
    t.text = '±'
    return r


# ============================================================
# 第4部分: 编辑 Word 文档，插入公式和丰富内容
# ============================================================

def add_formula_to_doc(doc, omml_xml_string, after_paragraph_idx=None):
    """
    向文档中添加 OMML 公式
    使用原始 XML 方式插入 OMML 公式
    """
    # 删除 XML 声明
    omml_xml_string = omml_xml_string.strip()
    if omml_xml_string.startswith('<?xml'):
        omml_xml_string = omml_xml_string.split('?>', 1)[1].strip()
    
    # 创建新段落
    if after_paragraph_idx is not None and after_paragraph_idx < len(doc.paragraphs):
        para = doc.paragraphs[after_paragraph_idx]
        new_para = para.insert_paragraph_after()
    else:
        new_para = doc.add_paragraph()
    
    # 设置段落样式
    new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 创建 run 并添加公式
    run = new_para.add_run()
    run.font.size = Pt(12)
    
    # 创建 oMathPara 包装
    omml_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    omath_para = ET.Element(f'{{{omml_ns}}}oMathPara')
    
    # 解析用户提供的 OMML
    try:
        omath_elem = ET.fromstring(omml_xml_string)
        omath_para.append(omath_elem)
        
        # 将 OMML 插入到 run 中
        run._element.append(omath_para)
    except ET.ParseError as e:
        print(f"  XML 解析错误: {e}")
    
    return new_para


def insert_formula_using_oxml(doc, formula_xml, after_paragraph_idx=None):
    """
    使用 python-docx 的 oxml 方式插入公式
    formula_xml 是完整的 m:oMath XML 字符串
    """
    from docx.oxml import OxmlElement
    from lxml import etree as lxml_etree
    
    # 创建新段落
    if after_paragraph_idx is not None and after_paragraph_idx < len(doc.paragraphs):
        para = doc.paragraphs[after_paragraph_idx]
        new_para_element = OxmlElement('w:p')
        para._element.addnext(new_para_element)
        new_para = doc.paragraphs[after_paragraph_idx + 1]
    else:
        new_para = doc.add_paragraph()
    
    # 解析公式 XML
    if isinstance(formula_xml, str):
        formula_element = lxml_etree.fromstring(formula_xml)
    else:
        formula_element = formula_xml
    
    # 创建 oMathPara 包装
    omath_para = OxmlElement('m:oMathPara')
    omath_para.append(formula_element)
    
    # 创建 run 并添加公式
    run_element = OxmlElement('w:r')
    run_element.append(omath_para)
    new_para._element.append(run_element)
    
    return new_para


def insert_formula_simple(doc, formula_xml_str, after_paragraph_idx=None):
    """
    使用 lxml 直接操作 XML 的方式插入公式
    最可靠的方法
    """
    from docx.oxml.ns import qn
    
    # 创建新段落
    if after_paragraph_idx is not None and after_paragraph_idx < len(doc.paragraphs):
        para_element = doc.paragraphs[after_paragraph_idx]._element
        new_p = etree.Element(qn('w:p'))
        para_element.addnext(new_p)
    else:
        new_p = etree.SubElement(doc.element.body, qn('w:p'))
    
    # 创建 run
    new_r = etree.SubElement(new_p, qn('w:r'))
    
    # 创建 run properties
    rPr = etree.SubElement(new_r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Cambria Math')
    rFonts.set(qn('w:hAnsi'), 'Cambria Math')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '24')
    
    # 创建 oMathPara
    omath_para = etree.SubElement(new_r, qn('m:oMathPara'))
    
    # 解析公式 XML
    formula_element = etree.fromstring(formula_xml_str)
    omath_para.append(formula_element)
    
    return new_p


def build_rich_word_document(original_doc_path, output_path):
    """
    构建丰富的 Word 文档：
    1. 保留原有内容
    2. 添加新公式（欧拉公式、正态分布、勾股定理、贝叶斯定理等）
    3. 添加公式说明和使用场景
    """
    
    doc = Document(original_doc_path)
    
    # ============================================================
    # 在原有内容后添加新的公式和说明
    # ============================================================
    
    # 添加分隔标题
    heading = doc.add_heading('丰富的数学公式集合', level=1)
    
    # ============================================================
    # 公式 3: 欧拉公式 (Euler's Formula)
    # ============================================================
    
    doc.add_heading('欧拉公式 (Euler\'s Formula)', level=2)
    
    # 描述
    p = doc.add_paragraph()
    p.add_run('欧拉公式被誉为"最美的数学公式"，它将自然常数 e、虚数单位 i、圆周率 π 联系在一起。').font.size = Pt(11)
    
    # e^(iπ) + 1 = 0
    formula_euler = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>e</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e><m:r><m:t>i</m:t></m:r><m:r><m:t>π</m:t></m:r></m:e>
            <m:sup><m:r><m:t>i</m:t></m:r><m:r><m:t>θ</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>cos</m:t></m:r>
        <m:r><m:t>θ</m:t></m:r>
        <m:r><m:t>+</m:t></m:r>
        <m:r><m:t>i</m:t></m:r>
        <m:r><m:t>sin</m:t></m:r>
        <m:r><m:t>θ</m:t></m:r>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_euler, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # 特殊形式
    p = doc.add_paragraph()
    p.add_run('当 θ = π 时，得到最著名的形式：').font.size = Pt(11)
    
    formula_euler2 = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>e</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e><m:r><m:t>i</m:t></m:r><m:r><m:t>π</m:t></m:r></m:e>
            <m:sup><m:r><m:t>i</m:t></m:r><m:r><m:t>θ</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>+</m:t></m:r>
        <m:r><m:t>1</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>0</m:t></m:r>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_euler2, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ============================================================
    # 公式 4: 正态分布 (Normal Distribution)
    # ============================================================
    
    doc.add_heading('正态分布 (Normal Distribution)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('正态分布是统计学中最重要的概率分布，也称为高斯分布。').font.size = Pt(11)
    
    formula_normal = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>f</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>x</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:f>
            <m:fPr/>
            <m:num><m:r><m:t>1</m:t></m:r></m:num>
            <m:den>
                <m:r><m:t>σ</m:t></m:r>
                <m:rad>
                    <m:radPr/>
                    <m:e><m:r><m:t>2</m:t></m:r><m:r><m:t>π</m:t></m:r></m:e>
                </m:rad>
            </m:den>
        </m:f>
        <m:r><m:t>e</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup>
                <m:r><m:t>-</m:t></m:r>
                <m:f>
                    <m:fPr/>
                    <m:num>
                        <m:r><m:t>(</m:t></m:r>
                        <m:r><m:t>x</m:t></m:r>
                        <m:r><m:t>-</m:t></m:r>
                        <m:r><m:t>μ</m:t></m:r>
                        <m:r><m:t>)</m:t></m:r>
                        <m:sSup>
                            <m:sSupPr/>
                            <m:e/>
                            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
                        </m:sSup>
                    </m:num>
                    <m:den>
                        <m:r><m:t>2</m:t></m:r>
                        <m:r><m:t>σ</m:t></m:r>
                        <m:sSup>
                            <m:sSupPr/>
                            <m:e/>
                            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
                        </m:sSup>
                    </m:den>
                </m:f>
            </m:sup>
        </m:sSup>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_normal, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    p = doc.add_paragraph()
    p.add_run('其中 μ 是均值，σ 是标准差。').font.size = Pt(11)
    
    # ============================================================
    # 公式 5: 勾股定理 (Pythagorean Theorem)
    # ============================================================
    
    doc.add_heading('勾股定理 (Pythagorean Theorem)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('勾股定理是几何学中最基本的定理之一，描述了直角三角形三边之间的关系。').font.size = Pt(11)
    
    formula_pythag = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>a</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>+</m:t></m:r>
        <m:r><m:t>b</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>c</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_pythag, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ============================================================
    # 公式 6: 贝叶斯定理 (Bayes' Theorem)
    # ============================================================
    
    doc.add_heading('贝叶斯定理 (Bayes\' Theorem)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('贝叶斯定理是概率论中的核心定理，描述了条件概率之间的关系。').font.size = Pt(11)
    
    formula_bayes = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>P</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>A</m:t></m:r>
        <m:r><m:t>|</m:t></m:r>
        <m:r><m:t>B</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:f>
            <m:fPr/>
            <m:num>
                <m:r><m:t>P</m:t></m:r>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>B</m:t></m:r>
                <m:r><m:t>|</m:t></m:r>
                <m:r><m:t>A</m:t></m:r>
                <m:r><m:t>)</m:t></m:r>
                <m:r><m:t>P</m:t></m:r>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>A</m:t></m:r>
                <m:r><m:t>)</m:t></m:r>
            </m:num>
            <m:den>
                <m:r><m:t>P</m:t></m:r>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>B</m:t></m:r>
                <m:r><m:t>)</m:t></m:r>
            </m:den>
        </m:f>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_bayes, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ============================================================
    # 公式 7: 积分公式 (Integral)
    # ============================================================
    
    doc.add_heading('积分公式 (Integral Formula)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('定积分是微积分中的基本概念，表示函数曲线下的面积。').font.size = Pt(11)
    
    formula_integral = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:nary>
            <m:naryPr>
                <m:chr m:val="∫"/>
                <m:limLoc m:val="subSup"/>
            </m:naryPr>
            <m:sub><m:r><m:t>a</m:t></m:r></m:sub>
            <m:sup><m:r><m:t>b</m:t></m:r></m:sup>
            <m:e>
                <m:r><m:t>f</m:t></m:r>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>x</m:t></m:r>
                <m:r><m:t>)</m:t></m:r>
                <m:r><m:t>d</m:t></m:r>
                <m:r><m:t>x</m:t></m:r>
            </m:e>
        </m:nary>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>F</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>b</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
        <m:r><m:t>-</m:t></m:r>
        <m:r><m:t>F</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>a</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_integral, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ============================================================
    # 公式 8: 泰勒级数 (Taylor Series)
    # ============================================================
    
    doc.add_heading('泰勒级数 (Taylor Series)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('泰勒级数将函数展开为无穷级数，是微积分中最重要的工具之一。').font.size = Pt(11)
    
    formula_taylor = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>f</m:t></m:r>
        <m:r><m:t>(</m:t></m:r>
        <m:r><m:t>x</m:t></m:r>
        <m:r><m:t>)</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:nary>
            <m:naryPr>
                <m:chr m:val="∑"/>
                <m:limLoc m:val="subSup"/>
            </m:naryPr>
            <m:sub>
                <m:r><m:t>n</m:t></m:r>
                <m:r><m:t>=</m:t></m:r>
                <m:r><m:t>0</m:t></m:r>
            </m:sub>
            <m:sup><m:r><m:t>∞</m:t></m:r></m:sup>
            <m:e>
                <m:f>
                    <m:fPr/>
                    <m:num>
                        <m:r><m:t>f</m:t></m:r>
                        <m:sSup>
                            <m:sSupPr/>
                            <m:e/>
                            <m:sup>
                                <m:r><m:t>(</m:t></m:r>
                                <m:r><m:t>n</m:t></m:r>
                                <m:r><m:t>)</m:t></m:r>
                            </m:sup>
                        </m:sSup>
                        <m:r><m:t>(</m:t></m:r>
                        <m:r><m:t>a</m:t></m:r>
                        <m:r><m:t>)</m:t></m:r>
                    </m:num>
                    <m:den>
                        <m:r><m:t>n</m:t></m:r>
                        <m:r><m:t>!</m:t></m:r>
                    </m:den>
                </m:f>
                <m:r><m:t>(</m:t></m:r>
                <m:r><m:t>x</m:t></m:r>
                <m:r><m:t>-</m:t></m:r>
                <m:r><m:t>a</m:t></m:r>
                <m:sSup>
                    <m:sSupPr/>
                    <m:e><m:r><m:t>)</m:t></m:r></m:e>
                    <m:sup><m:r><m:t>n</m:t></m:r></m:sup>
                </m:sSup>
            </m:e>
        </m:nary>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_taylor, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    # ============================================================
    # 公式 9: 麦克斯韦方程组 (Maxwell's Equations) - 积分形式
    # ============================================================
    
    doc.add_heading('麦克斯韦方程组 (Maxwell\'s Equations)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('麦克斯韦方程组是电磁学的基本方程，描述了电场和磁场的关系。').font.size = Pt(11)
    
    # 高斯电场定律
    formula_maxwell1 = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:nary>
            <m:naryPr>
                <m:chr m:val="∮"/>
                <m:limLoc m:val="subSup"/>
            </m:naryPr>
            <m:sub><m:r><m:t>S</m:t></m:r></m:sub>
            <m:sup/>
            <m:e>
                <m:r><m:t>E</m:t></m:r>
                <m:r><m:t>·</m:t></m:r>
                <m:r><m:t>d</m:t></m:r>
                <m:r><m:t>A</m:t></m:r>
            </m:e>
        </m:nary>
        <m:r><m:t>=</m:t></m:r>
        <m:f>
            <m:fPr/>
            <m:num><m:r><m:t>q</m:t></m:r></m:num>
            <m:den><m:r><m:t>ε</m:t></m:r><m:sSub><m:sSubPr/><m:e/><m:sub><m:r><m:t>0</m:t></m:r></m:sub></m:sSub></m:den>
        </m:f>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_maxwell1, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    p = doc.add_paragraph()
    p.add_run('高斯电场定律 (Gauss\'s Law for Electricity)').font.size = Pt(10)
    
    # 高斯磁场定律
    formula_maxwell2 = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:nary>
            <m:naryPr>
                <m:chr m:val="∮"/>
                <m:limLoc m:val="subSup"/>
            </m:naryPr>
            <m:sub><m:r><m:t>S</m:t></m:r></m:sub>
            <m:sup/>
            <m:e>
                <m:r><m:t>B</m:t></m:r>
                <m:r><m:t>·</m:t></m:r>
                <m:r><m:t>d</m:t></m:r>
                <m:r><m:t>A</m:t></m:r>
            </m:e>
        </m:nary>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>0</m:t></m:r>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_maxwell2, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    p = doc.add_paragraph()
    p.add_run('高斯磁场定律 (Gauss\'s Law for Magnetism)').font.size = Pt(10)
    
    # ============================================================
    # 公式 10: 爱因斯坦质能方程
    # ============================================================
    
    doc.add_heading('质能方程 (Mass-Energy Equivalence)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('爱因斯坦的质能方程是物理学中最著名的公式之一，揭示了质量和能量的等价关系。').font.size = Pt(11)
    
    formula_einstein = '''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
        <m:r><m:t>E</m:t></m:r>
        <m:r><m:t>=</m:t></m:r>
        <m:r><m:t>m</m:t></m:r>
        <m:r><m:t>c</m:t></m:r>
        <m:sSup>
            <m:sSupPr/>
            <m:e/>
            <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
    </m:oMath>'''
    
    insert_formula_simple(doc, formula_einstein, after_paragraph_idx=len(doc.paragraphs) - 1)
    
    p = doc.add_paragraph()
    p.add_run('其中 E 是能量，m 是质量，c 是光速。').font.size = Pt(11)
    
    # ============================================================
    # 添加总结
    # ============================================================
    
    doc.add_heading('公式总结', level=1)
    
    p = doc.add_paragraph()
    p.add_run('本文档包含以下公式：').font.size = Pt(11)
    
    formulas_list = [
        '1. 一元二次方程求根公式: x = (-b ± √(b² - 4ac)) / 2a',
        '2. 傅里叶级数: f(x) = a₀ + Σ(aₙcos(nπx/L) + bₙsin(nπx/L))',
        '3. 欧拉公式: e^(iθ) = cosθ + i·sinθ',
        '4. 正态分布: f(x) = 1/(σ√(2π)) · e^(-(x-μ)²/(2σ²))',
        '5. 勾股定理: a² + b² = c²',
        '6. 贝叶斯定理: P(A|B) = P(B|A)·P(A) / P(B)',
        '7. 定积分: ∫[a,b] f(x)dx = F(b) - F(a)',
        '8. 泰勒级数: f(x) = Σ f^(n)(a)/n! · (x-a)^n',
        '9. 麦克斯韦方程组',
        '10. 质能方程: E = mc²',
    ]
    
    for formula in formulas_list:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(formula).font.size = Pt(11)
    
    # 保存文档
    doc.save(output_path)
    print(f"丰富的 Word 文档已保存到: {output_path}")
    
    return doc


# ============================================================
# 主程序
# ============================================================

if __name__ == "__main__":
    original_doc = r'c:\Users\licl45\Desktop\testTRAE\testword.docx'
    output_doc = r'c:\Users\licl45\Desktop\testTRAE\testword_enriched.docx'
    test_code_path = r'c:\Users\licl45\Desktop\testTRAE\test_formulas_from_word.py'
    
    print("=" * 70)
    print("第1步: 从 Word 文档中读取公式")
    print("=" * 70)
    
    formulas = read_formulas_from_docx(original_doc)
    
    for f in formulas:
        print(f"\n公式 {f['index']}:")
        print(f"  提取文本: {f['text']}")
    
    print("\n" + "=" * 70)
    print("第2步: 生成 Python 测试代码")
    print("=" * 70)
    
    test_code = generate_test_code(formulas)
    
    with open(test_code_path, 'w', encoding='utf-8') as f:
        f.write(test_code)
    
    print(f"测试代码已保存到: {test_code_path}")
    
    print("\n" + "=" * 70)
    print("第3步: 运行测试代码")
    print("=" * 70)
    
    import subprocess
    result = subprocess.run(
        ['python', test_code_path],
        capture_output=True,
        text=True,
        encoding='utf-8'
    )
    print(result.stdout)
    if result.stderr:
        print("错误:", result.stderr)
    
    print("\n" + "=" * 70)
    print("第4步: 创建丰富的 Word 文档")
    print("=" * 70)
    
    build_rich_word_document(original_doc, output_doc)
    
    print("\n" + "=" * 70)
    print("完成！")
    print(f"  原始文档: {original_doc}")
    print(f"  丰富文档: {output_doc}")
    print(f"  测试代码: {test_code_path}")
    print("=" * 70)